import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def read_docx_content(file_path):
    """Legge tutto il testo del docx."""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text), doc

def extract_ped_page(doc):
    """
    Estrae il contenuto della prima pagina (o prima tabella) dove risiede il PED.
    Assume che il PED sia nei primi paragrafi o nella prima tabella.
    """
    ped_text = ""
    # Se il PED è in una tabella (molto comune nei brief)
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                ped_text += cell.text + " | "
            ped_text += "\n"
    
    # Aggiungi anche i primi paragrafi nel caso non sia in tabella
    # Limitiamo ai primi 1000 caratteri per sicurezza se non strutturato
    intro_text = "\n".join([p.text for p in doc.paragraphs[:10]])
    
    return f"DATI PED (TABELLA):\n{ped_text}\n\nCONTESTO INIZIALE:\n{intro_text}"

def highlight_text_in_docx(doc, text_snippets_to_highlight, output_path):
    """
    Cerca snippet di testo e li evidenzia in GIALLO.
    Nota: È una ricerca 'best effort' dato che la formattazione Word spezza le run.
    """
    for paragraph in doc.paragraphs:
        for snippet in text_snippets_to_highlight:
            if snippet in paragraph.text:
                # Metodo semplificato: evidenzia l'intero paragrafo se contiene lo snippet
                # Per evidenziare solo la parola serve iterare sulle 'runs', molto complesso
                # se la parola è spezzata tra stili diversi.
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    doc.save(output_path)

def save_text_as_docx(text, output_path, reference_doc_path=None):
    """Crea un nuovo docx dal testo corretto."""
    doc = docx.Document()
    if reference_doc_path:
        # Potremmo voler copiare stili o header, qui facciamo un dump semplice
        pass
    
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path)
